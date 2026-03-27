from __future__ import annotations

import json
import os
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from flask import Blueprint, current_app, jsonify, request
from openpyxl import load_workbook
from werkzeug.utils import secure_filename

mapper_bridge_bp = Blueprint("mapper_bridge_bp", __name__)

_DATA_DIR = Path(__file__).resolve().parent.parent / "cache" / "mapper_bridge"
_DATA_DIR.mkdir(parents=True, exist_ok=True)
_PROJECTS_PATH = _DATA_DIR / "projects.json"
if not _PROJECTS_PATH.exists():
    _PROJECTS_PATH.write_text("[]", encoding="utf-8")


def _timestamp() -> str:
    return datetime.utcnow().isoformat() + "Z"


def _load_projects() -> List[Dict[str, Any]]:
    try:
        return json.loads(_PROJECTS_PATH.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return []


def _save_projects(records: List[Dict[str, Any]]) -> None:
    _PROJECTS_PATH.write_text(json.dumps(records, indent=2), encoding="utf-8")


def _infer_field_type(label: str) -> str:
    lowered = label.lower()
    if any(key in lowered for key in ("date", "dob", "birth")):
        return "date"
    if any(key in lowered for key in ("amount", "total", "balance")):
        return "currency"
    if any(key in lowered for key in ("id", "number", "code")):
        return "identifier"
    if any(key in lowered for key in ("name", "owner", "person")):
        return "text"
    return "text"


def _is_required(label: str) -> bool:
    lowered = label.lower().strip()
    return "required" in lowered or lowered.endswith("*")


def _sanitize_heading(text: str) -> str:
    return text.strip().replace("\n", " ")


def _read_upload_bytes(field_name: str = "file") -> tuple[bytes, str]:
    upload = request.files.get(field_name)
    if not upload or not upload.filename:
        raise ValueError("file upload with form field 'file' is required")
    filename = secure_filename(upload.filename)
    data = upload.read()
    upload.close()
    return data, filename


def _analyze_word_template(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    document = Document(BytesIO(file_bytes))
    targets: List[Dict[str, Any]] = []
    for idx, paragraph in enumerate(document.paragraphs):
        content = paragraph.text.strip()
        if not content:
            continue
        style = paragraph.style.name if paragraph.style else "Paragraph"
        if style.lower().startswith("heading") or content.endswith(":"):
            targets.append(
                {
                    "heading": _sanitize_heading(content.rstrip(":")),
                    "location": f"paragraph:{idx + 1}",
                    "pattern": style,
                    "suggestedType": _infer_field_type(content),
                    "required": _is_required(content),
                }
            )
    warnings = []
    if not targets:
        warnings.append("No heading-like content detected. Confirm the template uses Word heading styles.")
    summary = {
        "fileName": filename,
        "analyzedAt": _timestamp(),
        "paragraphCount": len(document.paragraphs),
        "detectedTargets": len(targets),
    }
    return {"summary": summary, "targets": targets, "warnings": warnings}


def _analyze_excel_workbook(file_bytes: bytes, filename: str, user_context: Dict[str, Any] | None = None) -> Dict[str, Any]:
    workbook = load_workbook(BytesIO(file_bytes), data_only=True)
    sheet_summaries: List[Dict[str, Any]] = []
    total_columns = 0
    
    # Extract context hints if provided
    file_purpose = user_context.get("filePurpose") if user_context else None
    domain_context = user_context.get("domainContext", "") if user_context else ""
    focus_areas = user_context.get("focusAreas", []) if user_context else []
    
    for sheet in workbook.worksheets:
        header_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True), [])
        clean_headers = [value for value in header_row if value]
        total_columns += len(clean_headers)
        
        # Enhance sheet summary with context-aware descriptions
        sheet_info: Dict[str, Any] = {
            "sheetName": sheet.title,
            "rowCount": sheet.max_row,
            "columnCount": sheet.max_column,
            "sampleColumns": clean_headers[:10],
        }
        
        # Add context-enriched metadata if domain context provided
        if domain_context and clean_headers:
            sheet_info["contextNote"] = f"Interpreted as {file_purpose or 'data'} in {domain_context[:50]}..."
        
        sheet_summaries.append(sheet_info)
    
    summary = {
        "fileName": filename,
        "analyzedAt": _timestamp(),
        "sheetCount": len(workbook.sheetnames),
    }
    context = {
        "sheet_summary": sheet_summaries,
        "total_source_columns": total_columns,
    }
    warnings = []
    if not sheet_summaries:
        warnings.append("No worksheets found in the uploaded Excel file.")
    elif total_columns == 0:
        warnings.append("Header row appears empty on all sheets.")
    return {"summary": summary, "context": context, "warnings": warnings}


@mapper_bridge_bp.route("/api/health", methods=["GET"])
def mapper_health_check():
    version = os.getenv("SNOWCHAT_VERSION") or current_app.config.get("SNOWCHAT_VERSION", "dev")
    response = {
        "status": "ok",
        "message": "snowchat backend reachable",
        "timestamp": _timestamp(),
        "version": version,
    }
    return jsonify(response), 200


@mapper_bridge_bp.route("/mapping/parse/word", methods=["POST"])
def mapping_parse_word():
    try:
        file_bytes, filename = _read_upload_bytes()
        if not filename.lower().endswith(".docx"):
            return jsonify({"status": "error", "message": "only .docx files are supported"}), 400
        analysis = _analyze_word_template(file_bytes, filename)
        analysis["summary"]["receivedBytes"] = len(file_bytes)
        return jsonify(analysis), 200
    except ValueError as exc:
        return jsonify({"status": "error", "message": str(exc)}), 400
    except Exception as exc:  # noqa: BLE001 - catch and log for structured response
        current_app.logger.exception("Word template analysis failed")
        return jsonify({"status": "error", "message": "unable to analyze Word template", "details": str(exc)}), 500


@mapper_bridge_bp.route("/mapping/parse/excel", methods=["POST"])
def mapping_parse_excel():
    try:
        file_bytes, filename = _read_upload_bytes()
        if not filename.lower().endswith((".xlsx", ".xlsm")):
            return jsonify({"status": "error", "message": "only .xlsx or .xlsm files are supported"}), 400
        
        # Extract user context if provided
        context_str = request.form.get("context")
        user_context = None
        if context_str:
            try:
                user_context = json.loads(context_str)
                current_app.logger.info(f"Received user context: filePurpose={user_context.get('filePurpose')}, domain={user_context.get('domainContext')[:50] if user_context.get('domainContext') else 'none'}")
            except json.JSONDecodeError:
                current_app.logger.warning("Failed to parse context JSON, proceeding without context")
        
        analysis = _analyze_excel_workbook(file_bytes, filename, user_context)
        analysis["summary"]["receivedBytes"] = len(file_bytes)
        return jsonify(analysis), 200
    except ValueError as exc:
        return jsonify({"status": "error", "message": str(exc)}), 400
    except Exception as exc:  # noqa: BLE001
        current_app.logger.exception("Excel workbook analysis failed")
        return jsonify({"status": "error", "message": "unable to analyze Excel workbook", "details": str(exc)}), 500


@mapper_bridge_bp.route("/api/projects", methods=["GET"], strict_slashes=False)
@mapper_bridge_bp.route("/projects", methods=["GET"], strict_slashes=False)
def list_projects():
    projects = _load_projects()
    return jsonify({"status": "success", "data": projects, "timestamp": _timestamp()}), 200


@mapper_bridge_bp.route("/api/projects", methods=["POST"], strict_slashes=False)
@mapper_bridge_bp.route("/projects", methods=["POST"], strict_slashes=False)
def create_project():
    payload = request.get_json(silent=True) or {}
    name = (payload.get("name") or "").strip()
    if not name:
        return jsonify({"status": "error", "message": "Project name is required"}), 400
    owner = (payload.get("owner") or "system").strip() or "system"
    description = (payload.get("description") or "").strip()
    projects = _load_projects()
    record = {
        "id": f"proj-{len(projects) + 1}",
        "name": name,
        "description": description,
        "owner": owner,
        "createdAt": _timestamp(),
        "updatedAt": _timestamp(),
    }
    projects.append(record)
    _save_projects(projects)
    return jsonify({"status": "success", "message": "project created", "data": record}), 201


@mapper_bridge_bp.route("/api/ai/suggestions", methods=["POST"])
def ai_suggestions():
    payload = request.get_json(silent=True) or {}
    targets = payload.get("targets") or []
    project_name = payload.get("projectName") or payload.get("project") or "Untitled Project"
    suggestions: List[Dict[str, Any]] = []
    if not targets:
        suggestions.append(
            {
                "title": "Upload source templates",
                "confidence": 0.42,
                "details": "No targets detected yet. Upload Word or Excel templates to enable mapping simulation.",
                "generatedAt": _timestamp(),
            }
        )
    else:
        for target in targets:
            heading = target.get("heading") or target.get("name") or "Field"
            suggested_type = target.get("suggestedType") or "text"
            required = target.get("required", False)
            suggestions.append(
                {
                    "title": f"Review {heading}",
                    "confidence": 0.77 if required else 0.61,
                    "details": f"Detected {suggested_type} field. {'Mark as required' if required else 'Confirm optional status'} to lock the schema.",
                    "generatedAt": _timestamp(),
                }
            )
        suggestions.append(
            {
                "title": "Run coverage simulation",
                "confidence": 0.58,
                "details": f"{len(targets)} candidate targets detected for {project_name}. Run the simulator to validate downstream APIs.",
                "generatedAt": _timestamp(),
            }
        )
    return jsonify({"status": "success", "data": {"project": project_name, "suggestions": suggestions}, "timestamp": _timestamp()}), 200
